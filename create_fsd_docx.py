from docx import Document
from docx.shared import Pt

content = {
    'title': 'Functional Specification Document - My Tasks TODO List App',
    'version': '1.1.0',
    'date': 'May 27, 2026',
    'author': 'Lioe Inyan',
    'sections': [
        {
            'heading': 'Project Overview',
            'paragraphs': [
                'My Tasks is a mobile-friendly todo list web application built with React, TypeScript, Vite, and IndexedDB. The application enables users to manage tasks locally, without backend services, while delivering a polished, responsive user experience across phones, tablets, and desktops.',
                'The current version supports task categories, grouped task presentation, and category-based filtering in addition to the existing status filters.'
            ]
        },
        {
            'heading': 'Purpose',
            'paragraphs': [
                'The purpose of this project is to provide a lightweight, offline-capable task management application with a modern visual design and user-friendly interaction model. It should run in web browsers and remain usable on mobile devices without code changes.',
                'The category feature helps users organize tasks by domain, such as makanan, peralatan, pekerjaan, rumah, or any other free-text category.'
            ]
        },
        {
            'heading': 'Scope',
            'paragraphs': [
                'This document covers the functional requirements, application architecture, user interface design, data storage strategies, and user workflows for the My Tasks application.',
                'The scope includes task category creation through the add/edit task forms, grouped task display by category, and filtering the task list by a selected category.'
            ]
        },
        {
            'heading': 'Functional Requirements',
            'paragraphs': [
                '1. Create new tasks with a title.\n2. Optionally add task descriptions.\n3. Optionally assign a free-text category, for example makanan or peralatan.\n4. Optionally assign due dates.\n5. Optionally set task priority as low, medium, or high.\n6. Mark tasks as completed or active.\n7. Edit existing tasks, including title, description, category, and due date.\n8. Delete tasks.\n9. Filter task list by All, Active, or Completed.\n10. Filter task list by category using automatically detected categories from existing tasks.\n11. Group visible tasks by category in the task list.\n12. Display uncategorized tasks under an Uncategorized group.\n13. Display task statistics: active count, completed count, total count.\n14. Persist tasks locally in IndexedDB to survive page reloads and browser restarts.'
            ]
        },
        {
            'heading': 'Non-Functional Requirements',
            'paragraphs': [
                '1. The interface must be responsive and mobile-friendly.\n2. The application must be fast and visually polished.\n3. Data must persist locally using IndexedDB.\n4. The project should use React and TypeScript.\n5. The build system must use Vite.\n6. The application must compile successfully without errors.\n7. The UI must be accessible and support keyboard interactions.\n8. Category filtering and grouping must work without requiring a backend or fixed category table.'
            ]
        },
        {
            'heading': 'Architecture',
            'paragraphs': [
                'The application has the following primary architecture components: React components for UI, a service module for IndexedDB access, and CSS styling for responsive layout. The main React component, App.tsx, manages application state and coordinates data operations and presentation.',
                'App.tsx derives the available category list from the current Todo collection, stores the selected category filter, and passes the selected filter into TodoList. TodoList applies status and category filters, then groups visible tasks by category before rendering TodoItem components.'
            ]
        },
        {
            'heading': 'Component Structure',
            'paragraphs': [
                'App\n- AddTodo\n- Status Filters\n- Category Filter\n- TodoList\n  - Category Group Header\n  - TodoItem\n- Clear Completed Button\n- Footer'
            ]
        },
        {
            'heading': 'Data Model',
            'paragraphs': [
                'The application uses a Todo object with the following fields:\n- id: string\n- title: string\n- description: string\n- completed: boolean\n- category: string (optional)\n- dueDate: string (optional)\n- priority: low | medium | high\n- createdAt: number\n- updatedAt: number',
                'The category field is optional to preserve compatibility with existing tasks created before the category feature was added. Empty or missing category values are treated as Uncategorized in the UI.'
            ]
        },
        {
            'heading': 'IndexedDB Storage',
            'paragraphs': [
                'The app uses a local IndexedDB database named todolistDB with a single object store called todos. The object store uses id as its primary key and includes indexes for completed, priority, and createdAt. This model ensures task data persists without server dependency.',
                'The category field is stored as part of each Todo record. Category filtering is currently performed in application state after tasks are loaded from IndexedDB, so no additional IndexedDB index is required for the current feature size.'
            ]
        },
        {
            'heading': 'User Interface',
            'paragraphs': [
                'The UI is built with responsive CSS and uses visual hierarchy, gradient backgrounds, card-style containers, buttons, and compact mobile styling. The expanded add-task form, status filters, category filter, group headers, and action buttons are designed for both desktop and mobile finger-friendly use.',
                'Category values are displayed as compact badges on task items. The task list displays a group header for each visible category and a count for the number of tasks in that category.'
            ]
        },
        {
            'heading': 'Workflows',
            'paragraphs': [
                '1. Add Task: User enters a title and optional details, including category, then submits the form.\n2. Edit Task: User clicks edit, modifies task fields including category, and saves changes.\n3. Complete Task: User toggles the checkbox to mark task complete or active.\n4. Delete Task: User removes the task using the delete button.\n5. Filter by Status: User selects All, Active, or Completed to display relevant tasks.\n6. Filter by Category: User selects a category from the category dropdown to display tasks in that category.\n7. Group Tasks: The visible task list is automatically grouped by category.\n8. Clear Completed: User removes all completed tasks using the clear completed button.'
            ]
        },
        {
            'heading': 'Category Filtering and Grouping Rules',
            'paragraphs': [
                '1. Categories are free-text values entered by the user.\n2. The category dropdown is generated from non-empty category values in the current task list.\n3. Selecting All Categories displays all visible tasks grouped by category.\n4. Selecting a specific category displays only tasks matching that category.\n5. Status filtering and category filtering are combined.\n6. Tasks without category values are grouped under Uncategorized.\n7. If a selected category no longer exists after edits or deletes, the application falls back to All Categories.'
            ]
        },
        {
            'heading': 'Project Files',
            'paragraphs': [
                'Key files included in the project:\n- src/App.tsx\n- src/components/AddTodo.tsx\n- src/components/TodoList.tsx\n- src/components/TodoItem.tsx\n- src/services/indexedDB.ts\n- src/App.css\n- src/index.css\n- README.md\n- PROJECT_DOCUMENTATION.md\n- GITHUB_SETUP.md\n- CODEX.md'
            ]
        },
        {
            'heading': 'Deployment',
            'paragraphs': [
                'The application can be deployed as a static site using Vercel, Netlify, or GitHub Pages. The build output is generated in the dist folder with Vite.'
            ]
        },
        {
            'heading': 'Testing and Validation',
            'paragraphs': [
                'The application should be validated by running npm run lint and npm run build, then confirming a successful production build. Functional validation includes verifying task creation, editing, deletion, status filtering, category filtering, grouped category display, clear completed behavior, and persistence after page reloads.'
            ]
        },
        {
            'heading': 'Open Issues and Enhancements',
            'paragraphs': [
                'Potential future enhancements include: dark mode, recurring tasks, search capabilities, export/import functionality, cloud sync, drag and drop ordering, task notifications, and configurable category management.'
            ]
        }
    ]
}

doc = Document()
doc.add_heading(content['title'], 0)
meta = doc.add_paragraph()
meta.add_run(f"Version: {content['version']}\n").bold = True
meta.add_run(f"Date: {content['date']}\n").bold = True
meta.add_run(f"Author: {content['author']}\n").bold = True

for section in content['sections']:
    doc.add_heading(section['heading'], level=1)
    for para in section['paragraphs']:
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(6)

doc.save('FSD.docx')
print('Saved FSD.docx')
